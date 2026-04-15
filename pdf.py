# This script combines all docx exports into a single document and converts to PDF

from pathlib import Path
from docx import Document
from docx.oxml.shared import OxmlElement
from copy import deepcopy
import subprocess
import shutil

def combine_docx_files():
    """Combine all docx files from chapters/exports into a single document"""
    
    exports_dir = Path("chapters/exports")
    
    # List of docx files in order
    docx_files = [
        "01-intro.docx",
        "02-program.docx",
        "03-project.docx",
        "04-monitoring.docx",
        "05-ecivis.docx",
        "06-ceqa.docx",
        "07-resilience.docx",
        "08-funding.docx",
        "09-scale.docx",
        "10-grantmgmt.docx",
        "11-pathways.docx",
        "13-closeout.docx",
        "14-review.docx",
        "forestree.docx",
        "pwf.docx",
        "washoe.docx",
    ]
    
    print("Combining docx files...")
    
    # Load first document as master
    first_file = exports_dir / docx_files[0]
    master_doc = Document(first_file)
    
    # Clear all body content to start fresh
    for para in master_doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    for table in master_doc.tables:
        tbl = table._element
        tbl.getparent().remove(tbl)
    
    # Process each chapter
    for i, docx_file in enumerate(docx_files):
        file_path = exports_dir / docx_file
        
        if not file_path.exists():
            print(f"  Warning: {docx_file} not found")
            continue
        
        try:
            # Open the document
            source_doc = Document(file_path)
            
            # Copy all relationships (images, etc.)
            for rel_id, rel in source_doc.part.rels.items():
                if "image" in rel.reltype or "media" in rel.reltype:
                    try:
                        source_part = rel.target_part
                        master_doc.part.relate_to(source_part, rel.reltype)
                    except:
                        pass
            
            # Copy body elements
            for element in source_doc.element.body:
                # Deep copy to preserve structure
                new_element = deepcopy(element)
                master_doc.element.body.append(new_element)
            
            # Add page break between chapters (except after last)
            if i < len(docx_files) - 1:
                # Add paragraph with page break
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                pageBreak = OxmlElement('w:pageBreakBefore')
                pageBreak.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1')
                pPr.append(pageBreak)
                p.append(pPr)
                master_doc.element.body.append(p)
            
            print(f"  Added {docx_file}")
        except Exception as e:
            print(f"  Error adding {docx_file}: {e}")
    
    # Save combined document
    output_path = exports_dir / "fba-guide.docx"
    master_doc.save(output_path)
    print(f"\nCombined document saved to {output_path}")
    
    return output_path


def convert_to_pdf(docx_path):
    """Convert docx to PDF using available tools"""
    
    pdf_path = docx_path.with_suffix('.pdf')
    print(f"\nConverting to PDF...")
    
    # Try OnlyOffice
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            capture_output=True,
            timeout=120
        )
        if result.returncode == 0 and pdf_path.exists():
            print(f"PDF created: {pdf_path}")
            return pdf_path
    except FileNotFoundError:
        pass
    except Exception as e:
        pass
    
    # Try unoconv (universal office converter)
    try:
        result = subprocess.run(
            ["unoconv", "-f", "pdf", "-o", str(pdf_path), str(docx_path)],
            capture_output=True,
            timeout=120
        )
        if result.returncode == 0 and pdf_path.exists():
            print(f"PDF created: {pdf_path}")
            return pdf_path
    except FileNotFoundError:
        pass
    except Exception as e:
        pass
    
    print(f"Error: Could not convert to PDF. No conversion tool found.")
    print(f"Please install: unoconv")
    print(f"For Arch Linux: sudo pacman -S unoconv")
    return None
    
    return output_path

def convert_to_pdf_legacy(docx_path):
    """Deprecated function"""
    pass

if __name__ == "__main__":
    docx_path = combine_docx_files()
    if docx_path:
        convert_to_pdf(docx_path)